"""사업계획서 문서 파서.

지원 형식:
- .pdf  → pypdf 로 텍스트 추출
- .docx → python-docx
- .md / .txt → 그대로 읽기
- 첨부 없이 ``text`` 인자만 → 그대로 사용

표·차트 많은 PDF 는 텍스트 추출이 깨질 수 있어 사용자에게 안내가 필요.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass

import discord


# 첨부 크기 상한 (Discord 봇은 25MB 까지 받을 수 있지만, 사업계획서는 보통 5MB 이하)
_MAX_ATTACHMENT_BYTES = 20 * 1024 * 1024


@dataclass
class ParsedDocument:
    """파싱 결과 표준 형식."""

    raw_text: str
    word_count: int
    has_citations: bool        # '[출처:' 또는 '(출처:' 패턴 검사 결과
    numbers_with_sources: int  # 숫자 옆 각주(예: '12,3001)') 추정 개수
    numbers_total: int         # 본문 내 의미 있는 숫자 총 개수
    source: str                # "pdf" | "docx" | "md" | "txt" | "inline"


class DocumentParseError(Exception):
    """문서 파싱 실패."""


class DocumentParser:
    """비동기 진입점: 첨부와 텍스트 입력을 표준 ParsedDocument 로 변환."""

    async def parse(
        self,
        attachment: discord.Attachment | None = None,
        text: str | None = None,
    ) -> ParsedDocument:
        """첨부 또는 텍스트를 받아 ParsedDocument 로 변환.

        사용자가 둘 다 제공하면 첨부를 우선 사용 (보통 첨부가 본문).
        """
        if attachment is None and not text:
            raise DocumentParseError("문서가 없어요. 첨부하거나 텍스트로 입력해 주세요.")

        if attachment is not None:
            return await self._parse_attachment(attachment)

        return self._build(text or "", source="inline")

    # -----------------------------------------------------------------
    async def _parse_attachment(self, attachment: discord.Attachment) -> ParsedDocument:
        if attachment.size > _MAX_ATTACHMENT_BYTES:
            raise DocumentParseError("파일이 너무 커요 (20MB 이하만 지원). 본문만 추출해 주세요.")

        filename = (attachment.filename or "").lower()
        data = await attachment.read()

        if filename.endswith(".pdf"):
            return self._build(self._extract_pdf(data), source="pdf")
        if filename.endswith(".docx"):
            return self._build(self._extract_docx(data), source="docx")
        if filename.endswith(".md"):
            return self._build(data.decode("utf-8", errors="replace"), source="md")
        if filename.endswith(".txt"):
            return self._build(data.decode("utf-8", errors="replace"), source="txt")

        # 알 수 없는 확장자는 일단 텍스트 디코드 시도
        try:
            return self._build(data.decode("utf-8"), source="txt")
        except UnicodeDecodeError as exc:
            raise DocumentParseError(
                "지원하지 않는 형식이에요. PDF, DOCX, MD, TXT 만 가능합니다."
            ) from exc

    # -----------------------------------------------------------------
    @staticmethod
    def _extract_pdf(data: bytes) -> str:
        try:
            from pypdf import PdfReader  # type: ignore
        except ImportError as exc:
            raise DocumentParseError("PDF 파서가 설치되지 않았어요.") from exc

        try:
            reader = PdfReader(io.BytesIO(data))
            pages = []
            for page in reader.pages:
                # extract_text 가 None 또는 빈 문자열 반환 가능
                pages.append(page.extract_text() or "")
            text = "\n\n".join(pages).strip()
        except Exception as exc:  # noqa: BLE001
            raise DocumentParseError(
                "PDF 파싱에 실패했어요. 표·차트가 많으면 Markdown 으로 변환 후 다시 시도해 주세요."
            ) from exc

        if not text:
            raise DocumentParseError(
                "PDF 에서 텍스트를 추출하지 못했어요. 스캔본일 수 있으니 텍스트로 변환해 주세요."
            )
        return text

    @staticmethod
    def _extract_docx(data: bytes) -> str:
        try:
            from docx import Document  # type: ignore
        except ImportError as exc:
            raise DocumentParseError("DOCX 파서가 설치되지 않았어요.") from exc

        try:
            doc = Document(io.BytesIO(data))
        except Exception as exc:  # noqa: BLE001
            raise DocumentParseError("DOCX 파싱에 실패했어요.") from exc

        # 단락 + 표 셀 텍스트를 모두 수집
        chunks: list[str] = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                chunks.append(row_text)
        return "\n".join(c for c in chunks if c.strip())

    # -----------------------------------------------------------------
    @staticmethod
    def _build(text: str, source: str) -> ParsedDocument:
        cleaned = text.strip()
        words = len(re.findall(r"\S+", cleaned))

        # 숫자 추출 — 연도·페이지 번호도 포함되지만 휴리스틱이라 OK
        numbers = re.findall(r"\d[\d,\.]{0,15}", cleaned)
        # 출처 표기 — '[출처:', '(출처:', '※ 출처', 각주 번호 'N)' 등
        citation_pattern = re.compile(r"(\[출처\s*[:：]|\(출처\s*[:：]|※\s*출처)")
        has_citations = bool(citation_pattern.search(cleaned))

        # 숫자 옆 각주 표기를 가진 항목 (예: "12,300명1)") — 대략적 카운트
        sourced_numbers_pattern = re.compile(r"\d[\d,\.]{1,15}\s*[\)\]\^]?\s*\d?\s*\)")
        numbers_with_sources = len(sourced_numbers_pattern.findall(cleaned))

        return ParsedDocument(
            raw_text=cleaned,
            word_count=words,
            has_citations=has_citations,
            numbers_with_sources=numbers_with_sources,
            numbers_total=len(numbers),
            source=source,
        )
